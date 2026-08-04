"""
Sample: Table 1 — row height comparison.
Option A: auto row height (current, text-tight).
Option B: minimum row height ~0.30 in, text vertically centred.
Option C: minimum row height ~0.40 in, text vertically centred.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), "..", "references",
                   "CPIPM-TableStyle-Sample.docx")

NAVY  = RGBColor(0x1E, 0x3A, 0x5F)
BLACK = RGBColor(0x11, 0x18, 0x27)
GREY  = RGBColor(0x6B, 0x72, 0x80)

ROWS = [
    ("Target application",  "UDI labeling inline inspection — single station, single scan lane"),
    ("Scanner",             "Cognex DM-475V-LBL on desktop stand adapted for conveyor use"),
    ("Barcode type",        "GS1 DataMatrix only"),
    ("Trigger mode",        "Photosensor hardware item-presence (default); auto-cycle and manual also supported"),
    ("Indicator pole",      "3–5 segment light pole with per-color steady/flash states; audible alarm — TBD Engineering"),
    ("Conveyor control",    "Pusher divert on fail; line-stop behaviour TBD Engineering"),
    ("Timeline",            "Wk 1–3: Build & test  |  Wk 4: PE trial run  |  Wk 5: Harden  |  Wk 6: Go-live"),
    ("Physical product",    "Small folding cartons (~cell phone to paperback size, 1.25–1.75\" thick), "
                            "conveyed flat, narrow edge leading — virus and diagnostic test kits"),
]

def shd_cell(cell, hex_fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_fill)
    tcPr.append(shd)

def set_cell_margins(cell, top=36, bottom=36, left=108, right=108):
    """Tight uniform margins — only left/right are meaningful when row height is fixed."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

def set_cell_valign(cell, align="center"):
    """Vertical alignment within the cell: top | center | bottom."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)

def set_row_min_height(row, twips):
    """Set a minimum row height (atLeast). 1440 twips = 1 inch."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"),   str(twips))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)

def make_table(doc, min_row_height_twips, label=""):
    """
    min_row_height_twips=0  → auto (current style, no minimum)
    min_row_height_twips=N  → each row is at least N twips tall, text centred
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(label)
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = NAVY

    tbl = doc.add_table(rows=1 + len(ROWS), cols=2)
    tbl.style = "Table Grid"

    for row in tbl.rows:
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(4.5)

    # Header
    hdr = tbl.rows[0]
    if min_row_height_twips:
        set_row_min_height(hdr, min_row_height_twips)
    for i, txt in enumerate(["Parameter", "Value"]):
        cell = hdr.cells[i]
        cell.text = txt
        set_cell_margins(cell)
        set_cell_valign(cell, "center")
        shd_cell(cell, "1E3A5F")
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for ri, (param, value) in enumerate(ROWS):
        row = tbl.rows[ri + 1]
        if min_row_height_twips:
            set_row_min_height(row, min_row_height_twips)
        fill = "EBF8F7" if ri % 2 == 0 else "FFFFFF"
        for ci, txt in enumerate([param, value]):
            cell = row.cells[ci]
            cell.text = txt
            set_cell_margins(cell)
            set_cell_valign(cell, "center")
            shd_cell(cell, fill)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)
                    if ci == 0:
                        run.bold = True; run.font.color.rgb = NAVY
                    else:
                        run.font.color.rgb = BLACK

    doc.add_paragraph()

# ── Build ─────────────────────────────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin    = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

p = doc.add_paragraph()
r = p.add_run("CPIPM Table Style Sample — row height options")
r.bold = True; r.font.size = Pt(13); r.font.color.rgb = NAVY

p2 = doc.add_paragraph()
r2 = p2.add_run(
    "All three options use the same 9.5 pt font with tight left/right cell margins. "
    "The only difference is the minimum row height and vertical centring.")
r2.italic = True; r2.font.size = Pt(9); r2.font.color.rgb = GREY

# A: auto (matches original)                 ~210 twips natural height
make_table(doc, min_row_height_twips=0,
           label="Option A — auto height (original style, no minimum)")

# B: ~0.30 in minimum  =  432 twips
make_table(doc, min_row_height_twips=432,
           label="Option B — min row height 0.30 in, text centred  (≈ +100% on single-line rows)")

# C: ~0.40 in minimum  =  576 twips
make_table(doc, min_row_height_twips=576,
           label="Option C — min row height 0.40 in, text centred  (≈ +90% on single-line rows)")

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f"Saved: {OUT}")
