"""
CPIPM-Project-Outline-v1.0.docx → CPIPM-Project-Outline-v1.3.docx

v1.3 changes:
  1. Title font: 36 pt → 28 pt
  2. Title paragraph "CP Inline" → "VCCS Command Pilot™ Inline"
  3. First prose TM: first body use of "Command Pilot Inline" → "Command Pilot™ Inline"
  4. General terminology: remaining "CP Inline" → "Command Pilot Inline"
  5. Version string: 1.0 → 1.3, date → 4 August 2026
  6. Proportional row heights — WHOLE DOCUMENT:
       L = 1       → 0.25"      (360 twips)
       L >= 2      → (L×0.25 − 0.10)"  converted to twips
  7. Vertical-centre all table cells
  8. Caveat paragraph inserted after Section 6 heading (Operator Panel)
  9. Version history table: append v1.3 row
"""

import math, os, copy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = os.path.join(os.path.dirname(__file__), "..", "..", "CPIPM-Project-Outline-v1.0.docx")
OUT = os.path.join(os.path.dirname(__file__), "..", "references", "CPIPM-Project-Outline-v1.3.docx")

TWIPS_PER_INCH = 1440
CHARS_PER_LINE = {1: 90, 2: 75, 3: 60, 4: 45, 5: 35, 6: 25}
CHARS_DEFAULT  = 20

# ── Height formula ─────────────────────────────────────────────────────────────
def calc_height_twips(visual_rows):
    """L=1 → 0.25". L≥2 → (L×0.25 − 0.10)" × 1440."""
    if visual_rows <= 1:
        return round(0.25 * TWIPS_PER_INCH)          # 360
    return round((visual_rows * 0.25 - 0.10) * TWIPS_PER_INCH)

# ── Text replacements ──────────────────────────────────────────────────────────
REPLACEMENTS = [
    ("Version 1.0   |   2026-08-03   |   DRAFT",
     "Version 1.3   |   4 August 2026   |   DRAFT"),
    ("CP INLINE v1.0",   "Command Pilot Inline v1.0"),
    ("CP INLINE",        "Command Pilot Inline"),
    ("CP / Engineering", "VCCS/PIPS / Engineering"),
    ("CP Inline is",     "Command Pilot Inline is"),
    ("CP Inline reuses", "Command Pilot Inline reuses"),
    ("CP Inline v",      "Command Pilot Inline v"),
    ("CP Inline —",      "Command Pilot Inline —"),
    ("CP Inline,",       "Command Pilot Inline,"),
    ("CP Inline.",       "Command Pilot Inline."),
    ("CP Inline ",       "Command Pilot Inline "),
    ("CP Inline",        "Command Pilot Inline"),
]

def replace_text(text):
    for old, new in REPLACEMENTS:
        text = text.replace(old, new)
    return text

def fix_paragraph(para):
    """Apply replacements; handle title exact-match specially."""
    if not para.runs:
        return
    full = "".join(r.text for r in para.runs)
    if full.strip() == "CP Inline":                        # title paragraph
        para.runs[0].text = "VCCS Command Pilot\u2122 Inline"
        for r in para.runs[1:]: r.text = ""
        return
    new = replace_text(full)
    if new != full:
        para.runs[0].text = new
        for r in para.runs[1:]: r.text = ""

# ── Table helpers ──────────────────────────────────────────────────────────────
def cell_visual_rows(cell, cpl):
    texts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    if not texts:
        return 0
    return sum(max(1, math.ceil(len(t) / cpl)) for t in texts)

def row_max_vr(table_row, n_cols):
    cpl  = CHARS_PER_LINE.get(n_cols, CHARS_DEFAULT)
    seen = set()
    mx   = 0
    for cell in table_row.cells:
        tid = id(cell._tc)
        if tid in seen: continue
        seen.add(tid)
        vr = cell_visual_rows(cell, cpl)
        if vr > mx: mx = vr
    return mx

def set_row_height(row, twips):
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    for el in trPr.findall(qn("w:trHeight")): trPr.remove(el)
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   str(twips))
    trH.set(qn("w:hRule"), "atLeast")
    trPr.append(trH)

def set_cell_valign(cell, align="center"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for el in tcPr.findall(qn("w:vAlign")): tcPr.remove(el)
    vA = OxmlElement("w:vAlign")
    vA.set(qn("w:val"), align)
    tcPr.append(vA)

# ── Caveat paragraph insertion ─────────────────────────────────────────────────
CAVEAT_TEXT = (
    "\u26a0\u2002 PLACEHOLDER \u2014 The panel layout shown below is indicative only. "
    "Zone structure, element placement, labels, and interaction behaviour all require "
    "dedicated UI/UX design work before WPF implementation begins. Do not treat this "
    "wireframe as a final specification."
)

def insert_paragraph_after(ref_para, text, bold=True, italic=False,
                            color_hex="C00000", pt_size=10):
    """Insert a new paragraph immediately after ref_para in the document body."""
    new_p = OxmlElement("w:p")

    # paragraph properties — keep default style
    new_pPr = OxmlElement("w:pPr")
    new_p.append(new_pPr)

    # run
    new_r = OxmlElement("w:r")
    new_rPr = OxmlElement("w:rPr")

    if bold:
        b = OxmlElement("w:b"); new_rPr.append(b)
    if italic:
        i = OxmlElement("w:i"); new_rPr.append(i)
    if color_hex:
        col = OxmlElement("w:color")
        col.set(qn("w:val"), color_hex)
        new_rPr.append(col)
    if pt_size:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(pt_size * 2))   # half-points
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), str(pt_size * 2))
        new_rPr.append(sz); new_rPr.append(szCs)

    new_r.append(new_rPr)

    new_t = OxmlElement("w:t")
    new_t.text = text
    new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_r.append(new_t)
    new_p.append(new_r)

    ref_para._element.addnext(new_p)

# ── Version-history row ────────────────────────────────────────────────────────
def append_version_row(table, version, date, author, changes):
    new_tr = copy.deepcopy(table.rows[-1]._tr)
    table._tbl.append(new_tr)
    new_row = table.rows[-1]
    for cell, val in zip(new_row.cells, [version, date, author, changes]):
        for para in cell.paragraphs:
            for run in para.runs: run.text = ""
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        if p.runs: p.runs[0].text = val
        else:      p.add_run(val)
        set_cell_valign(cell, "center")

# ── Main ───────────────────────────────────────────────────────────────────────
doc = Document(SRC)

# 1. Title font size: 36 pt → 28 pt
title_para = doc.paragraphs[0]
for run in title_para.runs:
    run.font.size = Pt(28)

# 2. Body paragraphs: terminology + first TM
tm_done = False
op_panel_para = None

for idx, para in enumerate(doc.paragraphs):
    fix_paragraph(para)

    # Capture operator panel heading reference (after fix so text is stable)
    if ("6." in para.text and "Operator Panel" in para.text
            and para.style and para.style.name == "Heading 1"):
        op_panel_para = para

    # First prose TM (past title block)
    if not tm_done and idx > 5:
        for run in para.runs:
            if "Command Pilot Inline" in run.text:
                run.text = run.text.replace(
                    "Command Pilot Inline", "Command Pilot\u2122 Inline", 1)
                tm_done = True
                break

# 3. Insert caveat after operator panel heading
if op_panel_para:
    insert_paragraph_after(op_panel_para, CAVEAT_TEXT,
                            bold=True, italic=False, color_hex="C00000", pt_size=10)
else:
    print("WARNING: Operator Panel heading not found — caveat not inserted")

# 4. Tables: text + valign + proportional row heights
tables = doc.tables
for ti, table in enumerate(tables):
    n_cols  = len(table.columns)
    is_last = (ti == len(tables) - 1)

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                fix_paragraph(para)
            set_cell_valign(cell, "center")
        vr    = row_max_vr(row, n_cols)
        twips = calc_height_twips(vr)
        set_row_height(row, twips)

    if is_last:
        append_version_row(
            table,
            "1.3", "4 August 2026", "VCCS/PIPS / Engineering",
            "Title font 36 pt \u2192 28 pt; proportional row heights (L\u00d70.25\u2212 0.10 for L\u22652) "
            "applied whole document; operator panel placeholder caveat added; "
            "title TM and first body-copy TM fixed."
        )

# 5. Save
os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f"Saved: {OUT}")

# Spot-check: Table 1 heights
print("\nTable 1 spot-check:")
for ri, row in enumerate(doc.tables[0].rows):
    vr = row_max_vr(row, 2)
    tw = calc_height_twips(vr)
    print(f"  r{ri:02d}  vr={vr}  {tw}tw  {tw/TWIPS_PER_INCH:.3f}\"")
