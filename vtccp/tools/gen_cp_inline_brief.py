"""Generate CP-Inline-Engineering-Brief-v1.0.docx"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), "..", "references",
                   "CP-Inline-Engineering-Brief-v1.0.docx")

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x1E, 0x3A, 0x5F)
TEAL  = RGBColor(0x0D, 0x94, 0x88)
GREY  = RGBColor(0x6B, 0x72, 0x80)
BLACK = RGBColor(0x11, 0x18, 0x27)

def set_colour(run, rgb):
    run.font.color.rgb = rgb

def heading(text, level=1, colour=NAVY, space_before=10, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    set_colour(run, colour)
    return p

def body(text, italic=False, colour=BLACK, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = italic
    set_colour(run, colour)
    return p

def bullet(text, bold_prefix=None, colour=BLACK):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(10)
        set_colour(rb, colour)
    r = p.add_run(text)
    r.font.size = Pt(10)
    set_colour(r, colour)
    return p

def rule():
    """Thin horizontal rule via paragraph border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"),  "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0D9488")
    pBdr.append(bottom)
    pPr.append(pBdr)

# ─────────────────────────────────────────────────────────────────────────────
# HEADER BLOCK
# ─────────────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("VCCS Command Pilot™ Inline — Engineering Brief")
r.bold = True
r.font.size = Pt(15)
set_colour(r, NAVY)

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(1)
r2 = p2.add_run("Prospective project overview for engineering partners     |     Version 1.0  ·  3 August 2026     |     DRAFT — subject to revision")
r2.font.size = Pt(8.5)
r2.italic = True
set_colour(r2, GREY)

rule()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 1 — What is Command Pilot?
# ─────────────────────────────────────────────────────────────────────────────
heading("1.  What is VCCS Command Pilot\u2122?")

body(
    "Command Pilot is a Windows desktop application — built on WPF / .NET 8 — "
    "that VCCS/PIPS operates in its verification laboratory to assess barcode print "
    "quality for pharmaceutical, medical device, and other customers.  It accepts scan "
    "results from a Cognex DataMan fixed-mount reader, decodes GS1 linear and 2D "
    "payloads (and direct-part marks), grades them against ISO/IEC 15415 and related "
    "standards, logs each result to Excel and an event log, and presents the outcome "
    "to a lab operator.  The application runs on a bench-top verifier station; "
    "it is not a production-line application."
)

body(
    "Command Pilot Inline is the inline production variant.  The same proven engine handles "
    "decode and grading; a new operator-facing panel and a new I/O relay assembly are "
    "added to close the loop with the physical line."
)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 2 — Platform architecture
# ─────────────────────────────────────────────────────────────────────────────
heading("2.  Platform architecture (relevant assemblies)")

p = doc.add_paragraph(style="List Bullet")
p.paragraph_format.space_after = Pt(3)
rb = p.add_run("DeviceInterface — "); rb.bold = True; rb.font.size = Pt(10); set_colour(rb, NAVY)
r  = p.add_run("Cognex DataMan SDK integration (port 44444), raw DMCC command client "
               "(port 23), trigger management, image acquisition.  Handles all reader "
               "communication; the rest of the stack never touches hardware directly.")
r.font.size = Pt(10); set_colour(r, BLACK)

p = doc.add_paragraph(style="List Bullet")
p.paragraph_format.space_after = Pt(3)
rb = p.add_run("ExcelEngine — "); rb.bold = True; rb.font.size = Pt(10); set_colour(rb, NAVY)
r  = p.add_run("Structured verification-record writer.  Produces session Excel workbooks "
               "(.xlsx / .xls) and an append-only JSONL event log.  Schema is "
               "configurable; this writer would likely be used for inline operation, "
               "though other reporting options are also possible.")
r.font.size = Pt(10); set_colour(r, BLACK)

p = doc.add_paragraph(style="List Bullet")
p.paragraph_format.space_after = Pt(3)
rb = p.add_run("ConfigEngine — "); rb.bold = True; rb.font.size = Pt(10); set_colour(rb, NAVY)
r  = p.add_run("Operator-configurable parameters managed at multiple security levels. "
               "Routine job settings (job name, operator ID, roll number, etc.) are "
               "accessible to production operators; grading thresholds and other "
               "quality-critical parameters are controlled at a more secure level, "
               "separate from day-to-day job configuration.")
r.font.size = Pt(10); set_colour(r, BLACK)

p = doc.add_paragraph(style="List Bullet")
p.paragraph_format.space_after = Pt(3)
rb = p.add_run("InlineIo  (new) — "); rb.bold = True; rb.font.size = Pt(10); set_colour(rb, TEAL)
r  = p.add_run("Relay-board abstraction layer.  Exposes IRelayBoard (hardware-agnostic), "
               "MockRelayBoard (development/test), IndicatorPoleController, and "
               "ConveyorInterruptController.  Designed to accommodate indicator pole, "
               "pusher-divert, emergency-stop (e-stop), and other output channels as "
               "assignments are confirmed by Engineering.  Exact line-control behaviour TBD.")
r.font.size = Pt(10); set_colour(r, BLACK)

p = doc.add_paragraph(style="List Bullet")
p.paragraph_format.space_after = Pt(6)
rb = p.add_run("Command Pilot Inline Operator Panel  (new) — "); rb.bold = True; rb.font.size = Pt(10); set_colour(rb, TEAL)
r  = p.add_run("Full-screen WPF panel designed for production-floor use: grade result, "
               "indicator status, session counters, operator controls.  Layout TBD with "
               "customer validation.")
r.font.size = Pt(10); set_colour(r, BLACK)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3 — Application to CP Inline
# ─────────────────────────────────────────────────────────────────────────────
heading("3.  How this applies to the inline application")

body(
    "The physical target is small folding cartons — roughly cell-phone to paperback-book "
    "footprint, 1.25–1.75 in. thick — conveyed flat with the narrow edge leading.  These "
    "are virus and other diagnostic test kits.  A Cognex DM-475V-LBL (adapted from the "
    "customer's existing desktop stand) reads the GS1 DataMatrix label as each carton passes."
)

body("The decode-and-grade loop is unchanged from the existing platform:", space_after=2)

for step in [
    "Photosensor (or alternative trigger — TBD) signals carton presence.",
    "DeviceInterface fires a software or hardware trigger to the reader.",
    "Reader decodes and returns the GS1 DataMatrix result + ISO/IEC 15415 grade.",
    "ExcelEngine appends a verification record; image archived if required.",
    "Grade result is handed to InlineIo for indicator and line-control output.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(step); r.font.size = Pt(10); set_colour(r, BLACK)

body("", space_after=2)
body(
    "Indicator pole colour and line-control response are driven by the grade result.  "
    "The current working assumption is a multi-colour indicator pole and at least one "
    "pusher-divert output for failing cartons; whether a failing scan stops the line, "
    "diverts the carton while the line continues, or triggers some other response is "
    "a line-control integration question to be resolved with engineering.  The InlineIo "
    "layer is deliberately hardware-agnostic so these decisions do not require code changes."
)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 4 — Reuse estimate and new-build scope
# ─────────────────────────────────────────────────────────────────────────────
heading("4.  Reuse and new-build scope")

rows = [
    ("DeviceInterface (reader comms, trigger, image)",    "Reuse as-is",          "High"),
    ("ExcelEngine (Excel + JSONL reporting)",             "Reuse as-is",          "High"),
    ("ConfigEngine (operator config, PIN lock)",          "Reuse as-is",          "High"),
    ("InlineIo relay board layer",                        "New — stub complete",  "—"),
    ("Command Pilot Inline Operator Panel (WPF)",          "New build",            "—"),
    ("Line-control integration (divert, stop, etc.)",     "TBD Engineering",      "—"),
    ("Relay board hardware + wiring",                     "TBD Engineering",      "—"),
]

tbl = doc.add_table(rows=1 + len(rows), cols=3)
tbl.style = "Table Grid"

hdr = tbl.rows[0].cells
for i, h in enumerate(["Component", "Status", "Confidence"]):
    hdr[i].text = h
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    tc = hdr[i]._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "1E3A5F")
    tcPr.append(shd)

for ri, (comp, status, conf) in enumerate(rows):
    row = tbl.rows[ri + 1].cells
    row[0].text = comp
    row[1].text = status
    row[2].text = conf
    fill = "EBF4F4" if ri % 2 == 0 else "FFFFFF"
    for cell in row:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  fill)
        tcPr.append(shd)

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 5 — Key open items for engineering
# ─────────────────────────────────────────────────────────────────────────────
heading("5.  Key open items for engineering")

open_items = [
    ("Relay board",         "Model, channel count, USB/serial interface — drives IRelayBoard implementation."),
    ("Pusher divert",       "Output type (relay, PLC digital out), pulse vs. latched, timing relative to "
                            "carton position.  At least one divert channel assumed; exact behaviour TBD."),
    ("Line-control mode",   "Fail = divert only (line keeps running) vs. fail = divert + stop.  "
                            "Consecutive-fail threshold (if any) TBD.  Red indicator may be momentary "
                            "rather than latched — confirm with customer."),
    ("Indicator pole",      "Lamp count, colour mapping, steady vs. flash behaviour per grade band.  "
                            "Current mapping is a working assumption; all thresholds negotiable."),
    ("Trigger",             "Photosensor type and placement, end-of-carton detection (second sensor or "
                            "decode event), debounce timing."),
    ("Grade threshold",     "Customer's pass/fail grade boundary for this specific product and "
                            "regulatory context.  Default 1.5 assumed; confirm."),
    ("Image policy",        "Which frames to archive (all / fail-only / none), storage path, "
                            "DMST upload workflow for failures."),
    ("Reader mounting",     "Stand adaptation for conveyor height and carton size; working distance and "
                            "FOV verification against carton footprint."),
]

for item, detail in open_items:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    rb = p.add_run(f"{item}: "); rb.bold = True; rb.font.size = Pt(10); set_colour(rb, NAVY)
    r  = p.add_run(detail);      r.font.size = Pt(10);                  set_colour(r,  BLACK)

rule()

# Footer note
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
r = p.add_run(
    "This document describes the current working concept.  All line-control behaviours, "
    "channel assignments, and grade thresholds are provisional and subject to revision as "
    "engineering requirements are confirmed.  Contact: VCCS / Product Identification and "
    "Processing Systems, Inc."
)
r.font.size = Pt(8); r.italic = True; set_colour(r, GREY)

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f"Saved: {OUT}")
